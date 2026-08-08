"""
蒸馏 docx 文档：提取所有段落、样式、图片、表格、列表
输出：distill-output.txt（纯文本+结构标记）
     + distill-images/（提取的图片）
     + distill-report.md（蒸馏报告）
"""
import os
import sys
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from PIL import Image
import io
import json
import re

DOCX = r"d:\ai\latex\蒸馏区\知行读书 · 参赛作品说明书.docx"
OUT_DIR = r"d:\ai\latex\蒸馏区\distill-output"
IMG_DIR = os.path.join(OUT_DIR, "distill-images")
TXT_OUT = os.path.join(OUT_DIR, "distill-output.txt")
REPORT_OUT = os.path.join(OUT_DIR, "distill-report.md")
JSON_OUT = os.path.join(OUT_DIR, "distill-structure.json")

os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(IMG_DIR, exist_ok=True)

doc = Document(DOCX)
lines = []
structures = []  # 结构化数据
all_images = []  # 图片元信息

# ---------- 1. 提取所有图片 ----------
def extract_images():
    """提取 docx 内嵌图片，按出现顺序记录 paragraph index 和尺寸"""
    img_index = 0
    para_idx = 0
    for p in doc.paragraphs:
        # 检查段落中的图片
        for run in p.runs:
            # 通过 xml 查找 blip
            blips = run.element.findall(".//" + qn("a:blip"))
            for blip in blips:
                rId = blip.get(qn("r:embed"))
                if rId and rId in doc.part.rels:
                    target = doc.part.rels[rId].target_part
                    if target.content_type.startswith("image/"):
                        ext = target.content_type.split("/")[-1]
                        if ext == "jpeg":
                            ext = "jpg"
                        img_filename = f"img_{img_index:03d}.{ext}"
                        img_path = os.path.join(IMG_DIR, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(target.blob)
                        # 获取图片尺寸
                        try:
                            img = Image.open(io.BytesIO(target.blob))
                            w, h = img.size
                        except Exception:
                            w, h = 0, 0
                        all_images.append({
                            "index": img_index,
                            "filename": img_filename,
                            "path": img_path,
                            "width_px": w,
                            "height_px": h,
                            "paragraph_index": para_idx,
                            "nearby_text": p.text[:80]
                        })
                        img_index += 1
        para_idx += 1
    return img_index

# ---------- 2. 解析段落样式 ----------
def get_para_style_info(p):
    """提取段落的核心样式信息"""
    info = {
        "text": p.text,
        "style": p.style.name,
        "alignment": str(p.alignment) if p.alignment else "default",
    }
    # 第一个 run 的字体属性作为代表
    if p.runs:
        first_run = p.runs[0]
        font = first_run.font
        info["font_name"] = font.name
        info["font_size_pt"] = font.size.pt if font.size else None
        info["bold"] = font.bold
        info["italic"] = font.italic
        if font.color and font.color.rgb:
            info["color"] = str(font.color.rgb)
    # 段落级缩进
    pf = p.paragraph_format
    if pf.first_line_indent:
        info["first_line_indent_pt"] = pf.first_line_indent.pt
    if pf.left_indent:
        info["left_indent_pt"] = pf.left_indent.pt
    return info

# ---------- 3. 提取表格 ----------
def extract_tables():
    """提取所有表格内容"""
    tables = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows_data.append(cells)
        tables.append({
            "index": i,
            "row_count": len(table.rows),
            "col_count": len(table.columns) if table.rows else 0,
            "rows": rows_data
        })
    return tables

# ---------- 4. 主流程 ----------
img_count = extract_images()
tables = extract_tables()

# 输出文本结构
lines.append("=" * 80)
lines.append(f"文档：{DOCX}")
lines.append(f"段落总数：{len(doc.paragraphs)}")
lines.append(f"图片总数：{img_count}")
lines.append(f"表格总数：{len(tables)}")
lines.append("=" * 80)
lines.append("")

# 段落遍历
for i, p in enumerate(doc.paragraphs):
    info = get_para_style_info(p)
    text = p.text.strip()
    style = p.style.name
    # 标记标题
    prefix = ""
    if "Heading" in style or "标题" in style:
        prefix = f"### [{style}] "
    elif "Title" in style or "标题" in style:
        prefix = f"### [TITLE-{style}] "
    elif style == "Normal":
        prefix = ""

    # 标记此段附近是否有图片
    nearby_imgs = [im for im in all_images if im["paragraph_index"] == i]
    img_marker = ""
    if nearby_imgs:
        img_names = ", ".join(im["filename"] for im in nearby_imgs)
        img_marker = f"  [[IMG: {img_names}]]"

    font_info = ""
    if info.get("font_size_pt"):
        font_info = f"  {{font={info.get('font_name')}, size={info['font_size_pt']}pt, bold={info.get('bold')}}}"

    lines.append(f"[P{i:03d}] [{style}] {prefix}{text}{img_marker}{font_info}")

    structures.append({
        "paragraph_index": i,
        "style": style,
        "text": text,
        "first_line_indent_pt": info.get("first_line_indent_pt"),
        "font_size_pt": info.get("font_size_pt"),
        "bold": info.get("bold"),
        "font_name": info.get("font_name"),
        "nearby_images": [im["filename"] for im in nearby_imgs],
    })

# 表格内容
lines.append("")
lines.append("=" * 80)
lines.append(f"表格内容（共 {len(tables)} 个）")
lines.append("=" * 80)
for tbl in tables:
    lines.append(f"\n--- Table {tbl['index']} ({tbl['row_count']}x{tbl['col_count']}) ---")
    for row in tbl["rows"]:
        lines.append(" | ".join(row))

# 图片信息
lines.append("")
lines.append("=" * 80)
lines.append(f"图片列表（共 {img_count} 张）")
lines.append("=" * 80)
for im in all_images:
    lines.append(f"#{im['index']:03d} {im['filename']} ({im['width_px']}x{im['height_px']}px) @ P{im['paragraph_index']:03d}  附近文字: {im['nearby_text']}")

with open(TXT_OUT, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

with open(JSON_OUT, "w", encoding="utf-8") as f:
    json.dump({
        "paragraph_count": len(doc.paragraphs),
        "image_count": img_count,
        "table_count": len(tables),
        "images": all_images,
        "tables": tables,
        "paragraphs": structures,
    }, f, ensure_ascii=False, indent=2)

print(f"[OK] 文本输出: {TXT_OUT}")
print(f"[OK] JSON输出: {JSON_OUT}")
print(f"[OK] 图片目录: {IMG_DIR}")
print(f"[OK] 提取了 {img_count} 张图片, {len(tables)} 个表格, {len(doc.paragraphs)} 段落")
