# -*- coding: utf-8 -*-
"""
word_template_engine.py — 公文/学术 Word 模板引擎

用法：
    from word_template_engine import WordTemplateEngine
    engine = WordTemplateEngine('configs/fangzheng.json')
    engine.add_title('文档标题')
    engine.add_heading1('一、引言')
    engine.add_body('正文内容...')
    engine.add_table(headers, rows, caption='表1：xxx')
    engine.add_equation(omath_element, caption='式(1)：xxx')
    engine.save('output.docx')

设计原则：
    AI 只产出结构化数据（调用引擎 API），模板引擎负责保真渲染。
    切换字体方案只需换 JSON 配置文件。
"""

import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# OMML 命名空间
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


class WordTemplateEngine:
    """公文/学术 Word 模板引擎"""

    def __init__(self, config_path=None):
        """
        初始化引擎，加载字体配置。

        Args:
            config_path: JSON 配置文件路径。如果为 None，使用默认方正配置。
        """
        if config_path is None:
            config_path = Path(__file__).parent / 'configs' / 'fangzheng.json'
        else:
            config_path = Path(config_path)

        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)

        self.doc = Document()
        self._setup_page()
        self._configure_styles()

    # ============================================================
    # 内部初始化
    # ============================================================

    def _setup_page(self):
        """页面设置"""
        page = self.config.get('page', {})
        for s in self.doc.sections:
            s.page_width = Cm(page.get('width', 21.0))
            s.page_height = Cm(page.get('height', 29.7))
            s.top_margin = Cm(page.get('margin_top', 2.54))
            s.bottom_margin = Cm(page.get('margin_bottom', 2.54))
            s.left_margin = Cm(page.get('margin_left', 3.17))
            s.right_margin = Cm(page.get('margin_right', 3.17))

    def _configure_styles(self):
        """配置 Word 内置样式（Normal + Heading 1-4）"""
        fonts = self.config['fonts']
        styles_cfg = self.config['styles']

        # Normal
        self._set_style('Normal', fonts['body'], styles_cfg['body'])

        # Heading 1-4
        heading_map = {
            'Heading 1': ('heading1', fonts['heading1']),
            'Heading 2': ('heading2', fonts['heading2']),
            'Heading 3': ('heading3', fonts['heading3']),
            'Heading 4': ('heading4', fonts['heading4']),
        }
        for style_name, (cfg_key, font_name) in heading_map.items():
            self._set_style(style_name, font_name, styles_cfg[cfg_key])

    def _set_style(self, style_name, cn_font, cfg):
        """设置单个样式的字体、字号、颜色、间距"""
        style = self.doc.styles[style_name]
        en_font = self.config['fonts']['english']
        style.font.name = en_font
        style.font.size = Pt(cfg['size'])
        if 'color' in cfg:
            r, g, b = cfg['color']
            style.font.color.rgb = RGBColor(r, g, b)
        # 中文字体
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cn_font)
        # 间距
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(cfg.get('space_before', 0))
        pf.space_after = Pt(cfg.get('space_after', 0))

    # ============================================================
    # 底层辅助
    # ============================================================

    def _make_run(self, para, text, cn_font, size=12, bold=False, color=None):
        """在段落中添加 run 并设置中英文字体"""
        en_font = self.config['fonts']['english']
        run = para.add_run(text)
        run.font.name = en_font
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            r, g, b = color
            run.font.color.rgb = RGBColor(r, g, b)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cn_font)
        return run

    def _set_spacing(self, para, before=0, after=0):
        pf = para.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _set_indent(self, para, chars=2):
        pf = para.paragraph_format
        pf.first_line_indent = Pt(12 * chars)

    def _get_font(self, key):
        return self.config['fonts'][key]

    def _get_color(self, key):
        c = self.config['colors'].get(key)
        if c:
            return tuple(c)
        return None

    # ============================================================
    # 公开 API：文档元素
    # ============================================================

    def add_title(self, text, subtitle=None, author=None):
        """
        添加文档大标题（方正大标宋，居中）

        Args:
            text: 主标题
            subtitle: 副标题（可选）
            author: 作者信息（可选）
        """
        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, before=36, after=8)
        self._make_run(p, text, self._get_font('title'),
                       size=self.config['styles']['title']['size'],
                       color=self._get_color('dark'))

        # 副标题
        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_spacing(p, before=4, after=12)
            self._make_run(p, subtitle, self._get_font('heading1'),
                           size=self.config['styles']['heading1']['size'],
                           color=self._get_color('sub'))

        # 作者（居中楷体，四号 14pt，与 LaTeX \Large 对应）
        if author:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_spacing(p, before=4, after=12)
            self._make_run(p, author, self._get_font('heading3'), size=14)

    def add_abstract(self, abstract_text, keywords_text=None):
        """
        添加摘要和关键词（标签黑体，内容楷体）

        Args:
            abstract_text: 摘要内容
            keywords_text: 关键词内容（可选，分号分隔）
        """
        fonts = self.config['fonts']

        # 摘要
        p = self.doc.add_paragraph()
        self._set_spacing(p, before=4, after=4)
        self._set_indent(p, 2)
        self._make_run(p, '摘要：', fonts['heading2'], size=12, bold=True)
        self._add_inline_runs(p, abstract_text, fonts['heading3'], size=12)

        # 关键词
        if keywords_text:
            p = self.doc.add_paragraph()
            self._set_spacing(p, before=4, after=16)
            self._set_indent(p, 2)
            self._make_run(p, '关键词：', fonts['heading2'], size=12, bold=True)
            self._make_run(p, keywords_text, fonts['heading3'], size=12)

    def add_toc(self):
        """插入 Word 原生自动目录（TOC 域）"""
        # 目录标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, before=20, after=12)
        self._make_run(p, '目   录', self._get_font('heading1'),
                       size=self.config['styles']['heading1']['size'],
                       color=self._get_color('main'))

        # TOC 域
        p = self.doc.add_paragraph()
        self._set_spacing(p, before=4, after=4)
        run = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = r'TOC \o "1-4" \h \z \u'
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        ph = OxmlElement('w:t')
        ph.text = '右键此处选择"更新域"即可生成自动目录'
        ph_run = OxmlElement('w:r')
        ph_run.append(ph)
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_begin)
        run._element.append(instr)
        run._element.append(fld_sep)
        run._element.append(ph_run)
        run._element.append(fld_end)

        # 目录后分页
        p_break = self.doc.add_paragraph()
        run_break = p_break.add_run()
        import docx.enum.text
        run_break.add_break(docx.enum.text.WD_BREAK.PAGE)

    def add_heading1(self, text):
        """一级标题（Heading 1，TOC 可识别）"""
        p = self.doc.add_heading('', level=1)
        self._make_run(p, text, self._get_font('heading1'),
                       size=self.config['styles']['heading1']['size'],
                       color=self._get_color('main'))

    def add_heading2(self, text):
        """二级标题（Heading 2）"""
        p = self.doc.add_heading('', level=2)
        self._make_run(p, text, self._get_font('heading2'),
                       size=self.config['styles']['heading2']['size'],
                       color=self._get_color('main'))

    def add_heading3(self, text):
        """三级标题（Heading 3）"""
        p = self.doc.add_heading('', level=3)
        self._make_run(p, text, self._get_font('heading3'),
                       size=self.config['styles']['heading3']['size'],
                       color=self._get_color('sub'))

    def add_heading4(self, text):
        """四级标题（Heading 4）"""
        p = self.doc.add_heading('', level=4)
        self._make_run(p, text, self._get_font('heading4'),
                       size=self.config['styles']['heading4']['size'],
                       color=self._get_color('sub'))

    def _add_inline_runs(self, para, text, cn_font, size=12, bold=False,
                         color=None):
        """向段落添加 run，支持行内 **加粗** → 黑体片段"""
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                self._make_run(para, part[2:-2], self._get_font('heading2'),
                               size=size, bold=True, color=color)
            else:
                self._make_run(para, part, cn_font, size=size,
                               bold=bold, color=color)

    def add_body(self, text, indent=True):
        """
        正文段落（方正书宋，首行缩进，1倍行距）

        支持行内 **加粗** 标记 → 黑体片段（如 "这是**重点**内容"）。
        """
        p = self.doc.add_paragraph()
        self._set_spacing(p, before=3, after=3)
        if indent:
            self._set_indent(p, 2)
        self._add_inline_runs(p, text, self._get_font('body'), size=12)

    def add_annotation(self, text):
        """注释段落（仿宋，五号，灰色）"""
        p = self.doc.add_paragraph()
        self._set_spacing(p, before=4, after=4)
        self._set_indent(p, 2)
        self._make_run(p, text, self._get_font('annotation'),
                       size=10.5, color=self._get_color('gray'))

    def add_caption(self, text):
        """图表标题（楷体，五号，居中）"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, before=4, after=4)
        self._make_run(p, text, self._get_font('heading3'), size=10.5)

    def add_figure_placeholder(self, caption_text, hint='[图片占位区]', note=None):
        """
        图片占位 + 图注（注释在更下方）

        Args:
            caption_text: 图注文本
            hint: 占位提示文字
            note: 注释文本（可选，放在图注下方）
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, before=8, after=4)
        self._make_run(p, hint, self._get_font('heading3'), size=10.5,
                       color=self._get_color('gray'))
        self.add_caption(caption_text)
        if note:
            self.add_annotation(note)

    def add_picture(self, image_path, caption_text=None, note=None, width_cm=14.0):
        """
        插入真实图片 + 图注（图注在下方，注释更下方）。

        图片文件不存在时自动回退为占位框。

        Args:
            image_path: 图片路径（相对/绝对）
            caption_text: 图注文本（居中，楷体）
            note: 注释文本（可选，放在图注下方）
            width_cm: 图片宽度（厘米）
        """
        if not Path(image_path).exists():
            self.add_figure_placeholder(caption_text,
                                        hint=f'[图片缺失: {image_path}]',
                                        note=note)
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, before=8, after=4)
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        if caption_text:
            self.add_caption(caption_text)
        if note:
            self.add_annotation(note)

    def add_table(self, headers, rows, caption_text=None, note=None):
        """
        创建表格（标题在上方，注释在下方，表头黑体居中，内容水平居中）

        Args:
            headers: 表头列表
            rows: 数据行列表（每行为列表）
            caption_text: 表格标题（可选）
            note: 注释文本（可选，放在表格下方）
        """
        if caption_text:
            self.add_caption(caption_text)

        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头：黑体加粗居中
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._make_run(p, h, self._get_font('heading2'), size=10.5, bold=True)

        # 数据行：全部水平居中
        for r, row_data in enumerate(rows):
            for c, val in enumerate(row_data):
                cell = table.rows[r + 1].cells[c]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._make_run(p, val, self._get_font('body'), size=10.5)

        if note:
            self.add_annotation(note)

    def add_references(self, refs):
        """
        添加参考文献列表

        Args:
            refs: 参考文献文本列表（不含序号）
        """
        for i, ref in enumerate(refs, 1):
            p = self.doc.add_paragraph()
            self._set_spacing(p, before=2, after=2)
            self._make_run(p, f'[{i}] ', self._get_font('body'), size=10.5)
            self._make_run(p, ref, self._get_font('body'), size=10.5)

    # ============================================================
    # OMML 数学公式工具
    # ============================================================

    def add_equation(self, omath_element, caption_text=None):
        """
        插入行间数学公式（OMML 原生格式）

        Args:
            omath_element: lxml etree 元素（oMath）
            caption_text: 公式编号说明（可选）
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_spacing(p, before=8, after=4)
        p._element.append(omath_element)
        if caption_text:
            self.add_caption(caption_text)

    # ============================================================
    # OMML 构建静态方法
    # ============================================================

    @staticmethod
    def math_run(text):
        """创建 OMML run"""
        r = etree.Element('{%s}r' % M_NS)
        t = etree.SubElement(r, '{%s}t' % M_NS)
        t.text = text
        return r

    @staticmethod
    def math_sup(base_text, sup_text):
        """上标 base^sup"""
        sSup = etree.Element('{%s}sSup' % M_NS)
        e = etree.SubElement(sSup, '{%s}e' % M_NS)
        e.append(WordTemplateEngine.math_run(base_text))
        sup = etree.SubElement(sSup, '{%s}sup' % M_NS)
        sup.append(WordTemplateEngine.math_run(sup_text))
        return sSup

    @staticmethod
    def math_sub(base_text, sub_text):
        """下标 base_sub"""
        sSub = etree.Element('{%s}sSub' % M_NS)
        e = etree.SubElement(sSub, '{%s}e' % M_NS)
        e.append(WordTemplateEngine.math_run(base_text))
        sub = etree.SubElement(sSub, '{%s}sub' % M_NS)
        sub.append(WordTemplateEngine.math_run(sub_text))
        return sSub

    @staticmethod
    def math_frac(num_text, den_text):
        """分数 num/den"""
        f = etree.Element('{%s}f' % M_NS)
        num = etree.SubElement(f, '{%s}num' % M_NS)
        num.append(WordTemplateEngine.math_run(num_text))
        den = etree.SubElement(f, '{%s}den' % M_NS)
        den.append(WordTemplateEngine.math_run(den_text))
        return f

    @staticmethod
    def build_omath(*elements):
        """组合多个 OMML 元素为 oMath"""
        omath = etree.Element('{%s}oMath' % M_NS)
        for elem in elements:
            omath.append(elem)
        return omath

    # ============================================================
    # 保存
    # ============================================================

    def save(self, output_path, update_toc=True):
        """
        保存文档。

        Args:
            output_path: 输出路径
            update_toc: 是否用 Word COM 自动更新 TOC 目录（默认 True）
        """
        self.doc.save(output_path)
        print(f"[OK] {output_path}")
        name = self.config.get('name', '未知')
        print(f"  配置方案：{name}")

        if update_toc:
            self._update_toc_com(output_path)

    def _update_toc_com(self, output_path):
        """用 Word COM 自动更新 TOC 并保存"""
        try:
            import win32com.client
            import os
            abs_path = os.path.abspath(output_path)
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(abs_path)
            for toc in doc.TablesOfContents:
                toc.Update()
            # 用 SaveAs2 强制保存为 docx 格式，避免 OMML 冲突
            doc.SaveAs2(abs_path, FileFormat=12)  # 12 = wdFormatXMLDocument
            doc.Close()
            word.Quit()
            print(f"  [OK] TOC 已自动更新")
        except Exception as e:
            print(f"  [WARN] TOC 自动更新失败: {e}")
            print(f"  请在 Word 中手动右键目录区域 → '更新域'")
            # 确保 Word 进程被清理
            try:
                taskkill = os.system('taskkill /f /im WINWORD.EXE >nul 2>&1')
            except:
                pass
