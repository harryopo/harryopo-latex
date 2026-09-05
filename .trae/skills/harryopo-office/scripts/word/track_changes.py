#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
track_changes — Word 修订输出（AI 二稿带修订记录）
==================================================

在 .docx 上以"原生修订标记"应用一组修改，输出二稿：AI 的删除显示为 w:del（红线删除）、
新增显示为 w:ins（下划线插入），用户在 Word 里逐条接受/拒绝——与 redline.py
（对比用户修改版出红线稿）互补，共同构成改稿循环的双向留痕。

修订结构遵循 ECMA-376（与 Word/Python-Redlines 输出一致）：
  插入: <w:ins w:author=".." w:id="n" w:date=".."><w:r><w:t>新文本</w:t></w:r></w:ins>
  删除: <w:del w:author=".." w:id="n" w:date=".."><w:r><w:delText>旧文本</w:delText></w:r></w:del>

设计决策（P1 评估结论）：不引入 docx npm（Node 子进程）——修订 OOXML 结构是
公开标准且样本简单，python-docx + lxml 直构即可，与全 Python 栈一致。

用法：
  python track_changes.py 初稿.docx 二稿.docx --rev '[{"op":"replace","find":"A","replace":"B"}]' --author "AI"
  （或 from track_changes import apply_revisions）
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = '{%s}' % W_NS
XML_NS = 'http://www.w3.org/XML/1998/namespace'


def _el(tag: str, attrs: dict = None, text: str = None):
    e = etree.SubElement(etree.Element('root'), f'{W}{tag}')
    for k, v in (attrs or {}).items():
        # xml: 前缀属性需用 XML 命名空间（lxml 不接受字面 'xml:space'）
        qname = f'{{{XML_NS}}}{k[4:]}' if k.startswith('xml:') else f'{W}{k}'
        e.set(qname, str(v))
    if text is not None:
        e.text = text
    return e


class _Reviser:
    """维护唯一修订 id，输出带 author/date 的 w:ins / w:del 元素。"""

    def __init__(self, author: str):
        self.author = author
        self._id = 1000
        self.date = datetime.now().isoformat(timespec='seconds')

    def _next_id(self) -> int:
        self._id += 1
        return self._id

    def _run(self, text: str, del_text: bool = False):
        r = _el('r')
        if del_text:
            t = _el('delText', {'xml:space': 'preserve'}, text)
        else:
            t = _el('t', {'xml:space': 'preserve'}, text)
        r.append(t)
        return r

    def wrap_ins(self, text: str):
        ins = _el('ins', {'author': self.author, 'id': self._next_id(), 'date': self.date})
        ins.append(self._run(text))
        return ins

    def wrap_del(self, text: str):
        d = _el('del', {'author': self.author, 'id': self._next_id(), 'date': self.date})
        d.append(self._run(text, del_text=True))
        return d


def _split_run_for_replace(p, run, find: str, replace: str, rv: _Reviser) -> bool:
    """在单个 run 内把 find 替换为 replace，以修订形式（删旧 + 插新）。"""
    text = run.text or ''
    idx = text.find(find)
    if idx < 0:
        return False
    before, after = text[:idx], text[idx + len(find):]
    r_el = run._element
    parent = r_el.getparent()
    pos = list(parent).index(r_el)

    # 复制原 run 的 rPr 保持格式一致
    import copy
    rpr = r_el.find(f'{W}rPr')

    def _mk_run(text_val):
        nr = _el('r')
        if rpr is not None:
            nr.append(copy.deepcopy(rpr))
        nr.append(_el('t', {'xml:space': 'preserve'}, text_val))
        return nr

    inserts = []
    if before:
        inserts.append(_mk_run(before))
    inserts.append(rv.wrap_del(find))
    inserts.append(rv.wrap_ins(replace))
    if after:
        inserts.append(_mk_run(after))

    parent.remove(r_el)
    for off, el in enumerate(inserts):
        parent.insert(pos + off, el)
    return True


def _insert_paragraph_after(p, rv: _Reviser, text: str):
    """在段落后插入一个"整段为插入修订"的新段落。"""
    import copy
    new_p = copy.deepcopy(p._p)
    # 清空内容 runs，只保留 pPr
    for child in list(new_p):
        if child.tag != f'{W}pPr':
            new_p.remove(child)
    ins = rv.wrap_ins(text)
    new_p.append(ins)
    p._p.addnext(new_p)
    return new_p


def apply_revisions(docx_path: str, out_path: str, revisions: list,
                    author: str = 'AI Review') -> dict:
    """在 docx 上应用修订并输出二稿。

    revisions 支持：
      {'op': 'replace', 'find': 旧文本, 'replace': 新文本}   # 段内替换（find 须在同一 run 内）
      {'op': 'insert_after', 'anchor': 锚点段文本, 'text': 新段文本}
      {'op': 'delete', 'find': 要删除的文本}                  # 段内删除（同 replace 规则，replace=''）
    返回统计：{applied, skipped[{reason}], out}
    """
    rv = _Reviser(author)
    d = Document(docx_path)
    applied, skipped = 0, []

    for i, rev in enumerate(revisions):
        op = rev.get('op')
        if op == 'insert_after':
            anchor, text = rev.get('anchor', ''), rev.get('text', '')
            hit = next((p for p in d.paragraphs if anchor and anchor in p.text), None)
            if hit is None:
                skipped.append({'index': i, 'reason': f'锚点未找到: {anchor[:30]}'})
                continue
            _insert_paragraph_after(hit, rv, text)
            applied += 1
            continue

        find = rev.get('find', '')
        replace = '' if op == 'delete' else rev.get('replace', '')
        done = False
        for p in d.paragraphs:
            if find not in p.text:
                continue
            # 先试单 run 内命中；跨 run 场景 MVP 报跳过（与 docxtpl 占位符跨 run 同类限制）
            for run in p.runs:
                if find in (run.text or ''):
                    if _split_run_for_replace(p, run, find, replace, rv):
                        done = True
                    break
            if done:
                break
        if done:
            applied += 1
        else:
            reason = '未找到目标文本' if not find else '目标文本跨 run（MVP 不支持，请缩短 find 至单 run 内）'
            skipped.append({'index': i, 'reason': reason})

    d.save(out_path)
    return {'applied': applied, 'skipped': skipped, 'out': out_path}


def verify_revisions(docx_path: str) -> dict:
    """校验二稿含修订标记，返回统计。"""
    import zipfile
    xml = zipfile.ZipFile(docx_path).read('word/document.xml').decode('utf-8')
    return {'ins': xml.count('<w:ins '), 'del': xml.count('<w:del ')}


def main() -> int:
    ap = argparse.ArgumentParser(
        description='在 docx 上以原生修订标记应用修改，输出二稿（AI 改稿留痕）')
    ap.add_argument('original', help='初稿 docx')
    ap.add_argument('output', help='输出二稿路径')
    ap.add_argument('--rev', required=True,
                    help='修订 JSON 数组，如 [{"op":"replace","find":"A","replace":"B"}]')
    ap.add_argument('--author', default='AI Review', help='修订作者名')
    args = ap.parse_args()

    try:
        revisions = json.loads(args.rev)
    except json.JSONDecodeError as exc:
        print(f'[FATAL] --rev 不是合法 JSON: {exc}', file=sys.stderr)
        return 2

    result = apply_revisions(args.original, args.output, revisions, args.author)
    counts = verify_revisions(args.output)
    print(f"[OK] 二稿: {result['out']}  应用 {result['applied']} 条修订"
          f"（插入 {counts['ins']} / 删除 {counts['del']} 处标记），跳过 {len(result['skipped'])} 条")
    for s in result['skipped']:
        print(f"  [跳过] #{s['index']}: {s['reason']}")
    return 0 if result['applied'] else 1


if __name__ == '__main__':
    sys.exit(main())
