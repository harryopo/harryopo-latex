# -*- coding: utf-8 -*-
r"""
make_example_template.py — 生成 docxtpl 示例模板 template.docx

演示的占位符能力：
    1. 普通变量        {{ project_name }} / {{ version }} / {{ date }}
    2. 对象字段        {{ owner.name }} / {{ owner.phone }}
    3. 行内条件        {% if need_abstract %}...{% endif %}
    4. 表格行循环      {%tr for task in tasks %}...{%tr endfor %}
    5. 图片占位        {{ logo }}（data 中用 {"image":..., "width_mm":...}）

运行后生成：template.docx（模板）+ demo.png（示例图片）
"""

from pathlib import Path
from zlib import compress, crc32

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).parent


def make_png(path, width=320, height=200, rgb=(46, 125, 50)):
    """生成纯色 PNG（标准库实现，无 PIL 依赖）"""
    raw = b''.join(
        b'\x00' + bytes(rgb) * width for _ in range(height))
    def chunk(tag, data):
        c = tag + data
        return (len(data).to_bytes(4, 'big') + c
                + crc32(c).to_bytes(4, 'big'))
    png = (b'\x89PNG\r\n\x1a\n'
           + chunk(b'IHDR', (width.to_bytes(4, 'big')
                             + height.to_bytes(4, 'big')
                             + b'\x08\x02\x00\x00\x00'))
           + chunk(b'IDAT', compress(raw))
           + chunk(b'IEND', b''))
    path.write_bytes(png)
    print(f'  生成图片: {path}')


def build():
    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '方正书宋_GBK')

    # ===== 封面区 =====
    doc.add_paragraph('{{ logo }}')          # 图片占位
    doc.add_paragraph('项目申报书')
    doc.add_paragraph('项目名称：{{ project_name }}')
    doc.add_paragraph('项目版本：{{ version }}')
    doc.add_paragraph('申报日期：{{ date }}')

    # ===== 负责人（对象字段）=====
    doc.add_paragraph('')
    doc.add_paragraph('项目负责人：{{ owner.name }}')
    doc.add_paragraph('联系方式：{{ owner.phone }}')

    # ===== 行内条件 =====
    doc.add_paragraph('')
    doc.add_paragraph('{% if need_abstract %}本申报书包含项目摘要说明。{% endif %}')

    # ===== 循环表格 =====
    # docxtpl 约定：{%tr for %} 和 {%tr endfor %} 必须各自独占一行，
    # 中间的数据行（{{ task.xxx }}）在渲染时被循环复制
    doc.add_paragraph('')
    doc.add_paragraph('任务分工表：')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '序号'
    hdr[1].text = '任务'
    hdr[2].text = '负责人'
    hdr[3].text = '状态'

    row = table.add_row().cells
    row[0].text = '{%tr for task in tasks %}'
    row = table.add_row().cells
    row[0].text = '{{ loop.index }}'
    row[1].text = '{{ task.name }}'
    row[2].text = '{{ task.owner }}'
    row[3].text = '{{ task.status }}'
    row = table.add_row().cells
    row[0].text = '{%tr endfor %}'

    # ===== 结尾变量 =====
    doc.add_paragraph('')
    doc.add_paragraph('项目总结：{{ summary }}')

    out = HERE / 'template.docx'
    doc.save(str(out))
    print(f'已生成模板: {out}')


if __name__ == '__main__':
    make_png(HERE / 'demo.png')
    build()
